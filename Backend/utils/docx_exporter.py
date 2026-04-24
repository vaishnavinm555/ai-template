import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_docx_report(content_text, filename, title="AI Generated Academic Document"):
    """
    Converts plain text or markdown-style text into a professional Word (.docx) document.
    Handles headings, tables, bolding, and icons.
    """
    output_path = os.path.join("output", filename)
    if not os.path.exists("output"):
        os.makedirs("output")

    doc = Document()

    # Title
    if title:
        t = doc.add_heading(title, 0)
        t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Process content
    lines = content_text.split('\n')
    table_data = []
    is_table = False

    for line in lines:
        stripped_line = line.strip()

        # Handle Tables
        if stripped_line.startswith('|') and stripped_line.endswith('|'):
            if re.match(r'^\|[\s\-\:\!\|]+\|$', stripped_line):
                continue
            
            cells = [cell.strip() for cell in stripped_line.split('|') if cell.strip() or (cell == '' and stripped_line.count('|') > 1)]
            if cells:
                table_data.append(cells)
                is_table = True
            continue
        else:
            if is_table:
                if table_data:
                    rows = len(table_data)
                    cols = max(len(row) for row in table_data)
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for r in range(rows):
                        for c in range(len(table_data[r])):
                            table.cell(r, c).text = table_data[r][c]
                    doc.add_paragraph() # Spacer after table
                    table_data = []
                is_table = False

        if not stripped_line:
            continue

        # Horizontal Divider
        if stripped_line == '---' or stripped_line == '***':
            doc.add_paragraph("_" * 50)
            continue

        # Blockquote
        if stripped_line.startswith('>'):
            p = doc.add_paragraph(stripped_line.lstrip('> ').strip(), style='Quote')
            continue

        # Headings
        if stripped_line.startswith('# '):
            doc.add_heading(stripped_line[2:], level=0)
        elif stripped_line.startswith('## '):
            doc.add_heading(stripped_line[3:], level=1)
        elif stripped_line.startswith('### '):
            doc.add_heading(stripped_line[4:], level=2)
        else:
            p = doc.add_paragraph()
            # Simple bold/italic translation
            text = stripped_line
            # This is a very basic regex parser for bold/italic in one line
            # We'll process bold first, then italic
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)

    # Catch last table
    if is_table and table_data:
        rows = len(table_data)
        cols = max(len(row) for row in table_data)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r in range(rows):
            for c in range(len(table_data[r])):
                table.cell(r, c).text = table_data[r][c]

    doc.save(output_path)
    return output_path
